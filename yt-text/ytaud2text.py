import os
import glob
import subprocess
import sys
import whisper
import random
import string
from docx import Document
from docx.shared import Pt
from datetime import datetime

os.environ["PATH"] += r";C:\ffmpeg\bin"

YOUTUBE_URL = "https://www.youtube.com/watch?v=ZaPbP9DwBOE&t=300s"

def generate_random_id(length=4):
    characters = string.ascii_lowercase + string.digits
    return ''.join(random.choice(characters) for i in range(length))

def group_segments_into_paragraphs(segments, sentences_per_para=5):
    """Groups Whisper segments into readable paragraphs."""
    paragraphs = []
    current_para = ""
    sentence_count = 0

    for segment in segments:
        text = segment['text'].strip()
        current_para += text + " "
        
        # Count sentences by looking for ending punctuation
        if text.endswith(('.', '?', '!')):
            sentence_count += 1
        
        if sentence_count >= sentences_per_para:
            paragraphs.append(current_para.strip())
            current_para = ""
            sentence_count = 0
            
    # Add any remaining text
    if current_para:
        paragraphs.append(current_para.strip())
        
    return paragraphs

def save_to_docx(segments, url, base_name="YouTube_Transcript"):
    unique_id = generate_random_id()
    filename = f"{base_name}_{unique_id}.docx"
    
    doc = Document()
    doc.add_heading('YouTube Video Transcription', 0)
    
    # Metadata
    p = doc.add_paragraph()
    p.add_run('Source: ').bold = True
    p.add_run(url)
    p.add_run(f'\nGenerated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    doc.add_heading('Transcript', level=1)

    # --- UPDATED: Add grouped paragraphs ---
    paragraphs = group_segments_into_paragraphs(segments)
    
    for para_text in paragraphs:
        para = doc.add_paragraph(para_text)
        # Add spacing between paragraphs
        para.paragraph_format.space_after = Pt(12)

    # Styling
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    doc.save(filename)
    print(f"--- SUCCESS ---")
    print(f"Document saved as: {filename}")

# --- Processing Logic ---

print("Downloading audio...")
output_template = "audio.%(ext)s"
subprocess.run(
    [sys.executable, "-m", "yt_dlp", "-f", "bestaudio", "-o", output_template, YOUTUBE_URL],
    stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
)

audio_files = glob.glob("audio.*")
if not audio_files:
    raise FileNotFoundError("Audio file not found.")
audio_path = audio_files[0]

print("Loading Whisper model...")
model = whisper.load_model("base")

print("Transcribing (using FP32 for CPU compatibility)...")
# Note: passing fp16=False to avoid the warning you saw earlier
res = model.transcribe(audio_path, fp16=False)

# Pass the 'segments' list instead of just 'text'
save_to_docx(res["segments"], YOUTUBE_URL)

if os.path.exists(audio_path):
    os.remove(audio_path)